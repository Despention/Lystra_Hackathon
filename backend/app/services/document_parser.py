import re
from dataclasses import dataclass, field
from pathlib import Path

import chardet
import fitz  # pymupdf
from docx import Document as DocxDocument


@dataclass
class Section:
    title: str
    level: int
    content: str


@dataclass
class ParsedDocument:
    full_text: str
    sections: list[Section] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


HEADING_PATTERN = re.compile(
    r"^(\d+(?:\.\d+)*\.?)\s+(.+)$"
)


def _detect_sections(text: str) -> list[Section]:
    lines = text.split("\n")
    sections: list[Section] = []
    current_title = ""
    current_level = 0
    current_lines: list[str] = []

    for line in lines:
        stripped = line.strip()
        m = HEADING_PATTERN.match(stripped)
        if m and len(stripped) < 120:
            if current_title or current_lines:
                sections.append(Section(
                    title=current_title or "Введение",
                    level=current_level,
                    content="\n".join(current_lines).strip(),
                ))
            num = m.group(1)
            current_level = num.rstrip(".").count(".") + 1
            current_title = stripped
            current_lines = []
        else:
            current_lines.append(line)

    if current_title or current_lines:
        sections.append(Section(
            title=current_title or "Введение",
            level=current_level,
            content="\n".join(current_lines).strip(),
        ))

    return sections


def parse_pdf(filepath: str) -> ParsedDocument:
    doc = fitz.open(filepath)
    pages_text = []
    for page in doc:
        pages_text.append(page.get_text("text"))
    full_text = "\n".join(pages_text)
    doc.close()

    return ParsedDocument(
        full_text=full_text,
        sections=_detect_sections(full_text),
        metadata={
            "page_count": len(pages_text),
            "word_count": len(full_text.split()),
            "char_count": len(full_text),
        },
    )


def parse_docx(filepath: str) -> ParsedDocument:
    doc = DocxDocument(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    full_text = "\n".join(paragraphs)

    return ParsedDocument(
        full_text=full_text,
        sections=_detect_sections(full_text),
        metadata={
            "paragraph_count": len(doc.paragraphs),
            "word_count": len(full_text.split()),
            "char_count": len(full_text),
        },
    )


def parse_txt(filepath: str) -> ParsedDocument:
    path = Path(filepath)
    raw = path.read_bytes()
    detected = chardet.detect(raw)
    encoding = detected.get("encoding", "utf-8") or "utf-8"
    full_text = raw.decode(encoding, errors="replace")

    return ParsedDocument(
        full_text=full_text,
        sections=_detect_sections(full_text),
        metadata={
            "word_count": len(full_text.split()),
            "char_count": len(full_text),
        },
    )


def parse_text_input(text: str) -> ParsedDocument:
    return ParsedDocument(
        full_text=text,
        sections=_detect_sections(text),
        metadata={
            "word_count": len(text.split()),
            "char_count": len(text),
        },
    )


def parse_document(filepath: str) -> ParsedDocument:
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(filepath)
    elif ext in (".docx",):
        return parse_docx(filepath)
    elif ext in (".txt", ".text"):
        return parse_txt(filepath)
    else:
        return parse_txt(filepath)


def chunk_document_for_agent(
    document: ParsedDocument,
    max_chars: int = 4500,
    min_merge_chars: int = 500,
    overlap: int = 200,
) -> list[Section]:
    """Split a document into chunks small enough for a single LLM call.

    Strategy:
    1. If document has detected sections, use them as base units.
       - Merge consecutive small sections (<min_merge_chars) together.
       - Split sections larger than max_chars into sub-chunks.
    2. If no sections detected, split full_text into fixed-size windows with overlap.

    Each returned Section has a meaningful title for agent prompt context.
    """
    # Path 1: document has sections
    if document.sections:
        chunks: list[Section] = []
        buffer_title_parts: list[str] = []
        buffer_content: list[str] = []
        buffer_level = 1

        def _flush():
            if not buffer_content and not buffer_title_parts:
                return
            title = " / ".join(buffer_title_parts) if buffer_title_parts else "Фрагмент"
            content = "\n".join(buffer_content).strip()
            chunks.append(Section(title=title, level=buffer_level, content=content))

        for s in document.sections:
            content_len = len(s.content)

            # Case A: section itself exceeds max_chars → split it into windows
            if content_len > max_chars:
                _flush()
                buffer_title_parts = []
                buffer_content = []
                text = s.content
                start = 0
                part = 1
                total_parts = (content_len + max_chars - 1) // max_chars
                while start < content_len:
                    end = min(start + max_chars, content_len)
                    chunks.append(Section(
                        title=f"{s.title} (часть {part}/{total_parts})",
                        level=s.level,
                        content=text[start:end],
                    ))
                    part += 1
                    start = end - overlap if end < content_len else end
                continue

            # Case B: small section — accumulate into buffer
            current_buf_len = sum(len(c) for c in buffer_content)
            if current_buf_len + content_len > max_chars and buffer_content:
                _flush()
                buffer_title_parts = []
                buffer_content = []

            buffer_title_parts.append(s.title)
            buffer_content.append(s.content)
            buffer_level = s.level

            # If buffer is substantial, flush it
            if sum(len(c) for c in buffer_content) >= max_chars - min_merge_chars:
                _flush()
                buffer_title_parts = []
                buffer_content = []

        _flush()

        # Filter out empty chunks
        chunks = [c for c in chunks if c.content.strip()]
        if chunks:
            return chunks

    # Path 2: no sections (or all were empty) — window-split full_text
    text = document.full_text
    if not text:
        return []

    windows: list[Section] = []
    step = max(1, max_chars - overlap)
    total_parts = max(1, (len(text) + step - 1) // step)
    for i, start in enumerate(range(0, len(text), step)):
        end = min(start + max_chars, len(text))
        windows.append(Section(
            title=f"Фрагмент {i + 1}/{total_parts}",
            level=1,
            content=text[start:end],
        ))
        if end >= len(text):
            break
    return windows
